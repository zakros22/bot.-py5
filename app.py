from flask import Flask, render_template, request, send_file, jsonify
import os
import re
from werkzeug.utils import secure_filename
from docx import Document
import uuid

app = Flask(__name__)

# إعدادات رفع الملفات
UPLOAD_FOLDER = 'uploads'
DOWNLOAD_FOLDER = 'downloads'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 ميجابايت

# إنشاء المجلدات إذا لم تكن موجودة
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

# قاموس اللهجات مع قواعد التحويل
DIALECTS = {
    'fusha': {
        'name': 'الفصحى',
        'rules': {}
    },
    'iraqi': {
        'name': 'عراقي',
        'rules': {
            'ماذا': 'شكو',
            'كيف': 'شلون',
            'لماذا': 'ليش',
            'هذا': 'هذاي',
            'هذه': 'هذي',
            'أريد': 'أريد',
            'جيد': 'زين',
            'سأذهب': 'راح اروح',
            'أكل': 'آكل',
            'شرب': 'أشرب',
            'الآن': 'هسة',
            'كثير': 'هواية',
            'قليل': 'شوية'
        }
    },
    'egyptian': {
        'name': 'مصري',
        'rules': {
            'ماذا': 'إيه',
            'كيف': 'إزاي',
            'لماذا': 'ليه',
            'هذا': 'دا',
            'هذه': 'دي',
            'أريد': 'عايز',
            'جيد': 'كويّس',
            'سأذهب': 'هروح',
            'أكل': 'آكل',
            'شرب': 'أشرب',
            'الآن': 'دلوقتي',
            'كثير': 'أوي',
            'قليل': 'شوية'
        }
    },
    'gulf': {
        'name': 'خليجي',
        'rules': {
            'ماذا': 'شو',
            'كيف': 'كيف',
            'لماذا': 'ليش',
            'هذا': 'هذا',
            'هذه': 'هذي',
            'أريد': 'أبي',
            'جيد': 'زين',
            'سأذهب': 'بتروح',
            'أكل': 'آكل',
            'شرب': 'أشرب',
            'الآن': 'الحين',
            'كثير': 'وايد',
            'قليل': 'شوي'
        }
    },
    'syrian': {
        'name': 'شامي (سوريا)',
        'rules': {
            'ماذا': 'شو',
            'كيف': 'كيف',
            'لماذا': 'ليش',
            'هذا': 'هيدا',
            'هذه': 'هيدي',
            'أريد': 'بدي',
            'جيد': 'منيح',
            'سأذهب': 'رح روح',
            'أكل': 'آكل',
            'شرب': 'أشرب',
            'الآن': 'هلق',
            'كثير': 'كتير',
            'قليل': 'شوي'
        }
    }
}

def allowed_file(filename):
    """التحقق من صيغة الملف"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_to_dialect(text, dialect_code):
    """تحويل النص إلى اللهجة المختارة"""
    if dialect_code == 'fusha' or dialect_code not in DIALECTS:
        return text
    
    rules = DIALECTS[dialect_code]['rules']
    result = text
    
    # تطبيق قواعد التحويل
    for original, dialect_word in rules.items():
        result = result.replace(original, dialect_word)
        result = result.replace(original + ' ', dialect_word + ' ')
    
    return result

def process_txt_file(file_path, dialect_code):
    """معالجة ملف txt"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # تقسيم النص إلى جمل
    sentences = content.split('\n')
    translated_sentences = []
    
    for sentence in sentences:
        if sentence.strip():
            converted = convert_to_dialect(sentence, dialect_code)
            translated_sentences.append(converted)
        else:
            translated_sentences.append('')
    
    # حفظ الملف المترجم
    output_filename = f"translated_{uuid.uuid4().hex[:8]}.txt"
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], output_filename)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(translated_sentences))
    
    return output_filename

def process_docx_file(file_path, dialect_code):
    """معالجة ملف Word"""
    doc = Document(file_path)
    new_doc = Document()
    
    for paragraph in doc.paragraphs:
        converted = convert_to_dialect(paragraph.text, dialect_code)
        new_doc.add_paragraph(converted)
    
    output_filename = f"translated_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(app.config['DOWNLOAD_FOLDER'], output_filename)
    new_doc.save(output_path)
    
    return output_filename

@app.route('/')
def index():
    """الصفحة الرئيسية"""
    return render_template('index.html', dialects=DIALECTS)

@app.route('/upload', methods=['POST'])
def upload_file():
    """رفع الملف وترجمته"""
    # التحقق من وجود ملف
    if 'file' not in request.files:
        return jsonify({'error': 'لا يوجد ملف'}), 400
    
    file = request.files['file']
    dialect = request.form.get('dialect', 'fusha')
    
    # التحقق من اختيار ملف
    if file.filename == '':
        return jsonify({'error': 'لم يتم اختيار ملف'}), 400
    
    # التحقق من صيغة الملف
    if not allowed_file(file.filename):
        return jsonify({'error': 'نوع الملف غير مدعوم. الأنواع المدعومة: txt, docx'}), 400
    
    try:
        # حفظ الملف المرفوع
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # معالجة الملف حسب نوعه
        ext = filename.rsplit('.', 1)[1].lower()
        
        if ext == 'txt':
            output_filename = process_txt_file(file_path, dialect)
        elif ext == 'docx':
            output_filename = process_docx_file(file_path, dialect)
        else:
            return jsonify({'error': 'نوع ملف غير مدعوم'}), 400
        
        # تنظيف الملف المرفوع
        os.remove(file_path)
        
        return jsonify({
            'success': True,
            'message': 'تمت الترجمة بنجاح!',
            'download_url': f'/download/{output_filename}'
        })
        
    except Exception as e:
        return jsonify({'error': f'حدث خطأ: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """تحميل الملف المترجم"""
    file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return jsonify({'error': 'الملف غير موجود'}), 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=False)
